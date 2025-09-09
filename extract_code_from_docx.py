#!/usr/bin/env python3
"""
Extract images from .docx files and use OCR to extract Python code.
Keep original images and append extracted code below each image.
"""

import os
import sys
import re
import zipfile
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import easyocr
import pytesseract

def setup_ocr():
    """Initialize OCR readers."""
    print("Initializing OCR engines...")
    # Initialize EasyOCR (better for general text)
    easy_reader = easyocr.Reader(['en'])
    return easy_reader

def extract_images_from_docx(docx_path, temp_dir):
    """Extract all images from a .docx file."""
    images = []
    
    # .docx files are zip archives
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        # Look for images in word/media/ directory
        for file_info in docx_zip.filelist:
            if file_info.filename.startswith('word/media/') and any(
                file_info.filename.lower().endswith(ext) 
                for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
            ):
                # Extract image to temp directory
                image_name = os.path.basename(file_info.filename)
                temp_image_path = os.path.join(temp_dir, image_name)
                
                with docx_zip.open(file_info.filename) as source_file:
                    with open(temp_image_path, 'wb') as target_file:
                        target_file.write(source_file.read())
                
                images.append(temp_image_path)
    
    return images

def is_code_like(text):
    """Check if text looks like Python code."""
    if not text.strip():
        return False
    
    # Common Python keywords and patterns
    python_indicators = [
        'def ', 'class ', 'import ', 'from ', 'if ', 'else:', 'elif ',
        'for ', 'while ', 'try:', 'except:', 'finally:', 'with ',
        'return ', 'yield ', 'lambda ', 'async ', 'await ',
        '__init__', '__main__', 'self.', 'print(', 'len(',
        '==', '!=', '<=', '>=', '->', '=>',
        # Indentation patterns
        '\n    ', '\n        ',
        # Common imports
        'pandas', 'numpy', 'matplotlib', 'sklearn', 'torch',
        'fastapi', 'flask', 'django', 'requests'
    ]
    
    # Check for multiple indicators
    indicator_count = sum(1 for indicator in python_indicators if indicator in text.lower())
    
    # Also check for structural patterns
    has_indentation = '    ' in text or '\t' in text
    has_brackets = '(' in text and ')' in text
    has_colons = ':' in text
    
    return indicator_count >= 2 or (indicator_count >= 1 and has_indentation and has_brackets)

def clean_code_text(text):
    """Clean up OCR text to make it more code-like."""
    if not text:
        return ""
    
    # Remove extra spaces and fix common OCR errors
    text = re.sub(r' +', ' ', text)  # Multiple spaces to single space
    text = re.sub(r'(?<=[a-zA-Z])\s+(?=\()', '', text)  # Remove space before parentheses
    text = re.sub(r'(?<=\()\s+', '', text)  # Remove space after opening parenthesis
    text = re.sub(r'\s+(?=\))', '', text)  # Remove space before closing parenthesis
    text = re.sub(r'(?<=:)\s*\n\s*', ':\n    ', text)  # Fix indentation after colons
    
    # Fix common OCR character mistakes in code
    replacements = {
        'O': '0',  # Letter O to zero (in specific contexts)
        'l': '1',  # Lowercase L to 1 (in specific contexts)
        '|': 'I',  # Pipe to capital I
        '§': 'S',  # Section symbol to S
    }
    
    # Apply replacements cautiously
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Preserve leading whitespace for indentation
        leading_space = len(line) - len(line.lstrip())
        content = line.strip()
        
        if content:
            cleaned_lines.append(' ' * leading_space + content)
        else:
            cleaned_lines.append('')
    
    return '\n'.join(cleaned_lines)

def ocr_image(image_path, easy_reader):
    """Extract text from image using OCR."""
    try:
        # Load image
        image = Image.open(image_path)
        
        # Try EasyOCR first (often better for code)
        try:
            easy_results = easy_reader.readtext(image_path, detail=0)
            easy_text = ' '.join(easy_results)
        except Exception as e:
            print(f"EasyOCR failed for {image_path}: {e}")
            easy_text = ""
        
        # Try Tesseract as backup/comparison
        try:
            tess_text = pytesseract.image_to_string(image, config='--psm 6')
        except Exception as e:
            print(f"Tesseract failed for {image_path}: {e}")
            tess_text = ""
        
        # Choose the better result (longer text that looks more like code)
        if is_code_like(easy_text) and len(easy_text) > len(tess_text) * 0.8:
            extracted_text = easy_text
            ocr_method = "EasyOCR"
        elif is_code_like(tess_text):
            extracted_text = tess_text
            ocr_method = "Tesseract"
        elif len(easy_text) > len(tess_text):
            extracted_text = easy_text
            ocr_method = "EasyOCR"
        else:
            extracted_text = tess_text
            ocr_method = "Tesseract"
        
        # Clean up the extracted text
        cleaned_text = clean_code_text(extracted_text)
        
        return cleaned_text, ocr_method
        
    except Exception as e:
        print(f"Error processing image {image_path}: {e}")
        return "", "Error"

def add_code_block_to_docx(doc, code_text, ocr_method, image_name):
    """Add extracted code block to document."""
    # Add a paragraph with extracted code header
    header_p = doc.add_paragraph()
    header_run = header_p.add_run(f"Extracted Python Code (from {image_name} using {ocr_method}):")
    header_run.bold = True
    
    # Add the code block
    code_p = doc.add_paragraph()
    code_run = code_p.add_run(code_text)
    
    # Try to format as code (monospace font)
    try:
        code_run.font.name = 'Courier New'
        code_run.font.size = Inches(0.15)  # Small font size
    except:
        pass  # If formatting fails, just use default
    
    # Add some spacing
    doc.add_paragraph()

def process_docx_file(input_path, output_path, easy_reader):
    """Process a single .docx file."""
    print(f"Processing: {input_path}")
    
    # Create temporary directory for images
    temp_dir = "temp_images"
    os.makedirs(temp_dir, exist_ok=True)
    
    try:
        # Extract images from the .docx file
        images = extract_images_from_docx(input_path, temp_dir)
        
        if not images:
            # No images found, just copy the file
            shutil.copy2(input_path, output_path)
            print(f"  No images found, copied as-is")
            return
        
        print(f"  Found {len(images)} images")
        
        # Load the original document
        doc = Document(input_path)
        
        # Process each image
        for i, image_path in enumerate(images):
            print(f"  Processing image {i+1}/{len(images)}: {os.path.basename(image_path)}")
            
            # Extract text using OCR
            extracted_code, ocr_method = ocr_image(image_path, easy_reader)
            
            # Only add code block if we extracted meaningful text
            if extracted_code.strip() and (is_code_like(extracted_code) or len(extracted_code) > 20):
                add_code_block_to_docx(doc, extracted_code, ocr_method, os.path.basename(image_path))
                print(f"    Extracted {len(extracted_code)} characters of potential code")
            else:
                print(f"    No meaningful code detected in image")
        
        # Save the modified document
        doc.save(output_path)
        print(f"  Saved to: {output_path}")
        
    except Exception as e:
        print(f"Error processing {input_path}: {e}")
        # Copy original file if processing fails
        shutil.copy2(input_path, output_path)
        
    finally:
        # Clean up temporary directory
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

def main():
    """Main processing function."""
    book_dir = Path("book")
    output_dir = Path("book_1")
    
    if not book_dir.exists():
        print("Error: 'book' directory not found!")
        return
    
    # Create output directory
    output_dir.mkdir(exist_ok=True)
    
    # Initialize OCR
    easy_reader = setup_ocr()
    
    # Find all .docx files
    docx_files = list(book_dir.glob("*.docx"))
    
    if not docx_files:
        print("No .docx files found in 'book' directory!")
        return
    
    print(f"Found {len(docx_files)} .docx files to process")
    
    # Process each file
    for docx_file in docx_files:
        output_file = output_dir / docx_file.name
        process_docx_file(docx_file, output_file, easy_reader)
    
    print(f"\nProcessing complete! Check the 'book_1' directory for results.")

if __name__ == "__main__":
    main()