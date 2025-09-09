#!/usr/bin/env python3
"""
Final script to process all .docx files and extract any text from images,
adding the extracted content below the original images.
"""

import os
import zipfile
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageEnhance
import pytesseract

def preprocess_image_for_ocr(image_path, output_path=None):
    """Preprocess image to improve OCR accuracy."""
    try:
        with Image.open(image_path) as img:
            # Convert RGBA to RGB with white background
            if img.mode == 'RGBA':
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize if too small for good OCR
            width, height = img.size
            if width < 800 or height < 600:
                scale = max(800/width, 600/height)
                new_width = int(width * scale)
                new_height = int(height * scale)
                img = img.resize((new_width, new_height), Image.LANCZOS)
            
            # Enhance contrast and sharpness
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.2)
            
            # Convert to grayscale for better OCR
            img = img.convert('L')
            
            if output_path:
                img.save(output_path)
            
            return img
    except Exception as e:
        print(f"    Error preprocessing image: {e}")
        return None

def extract_text_with_ocr(image_path):
    """Extract text from image using OCR with multiple configurations."""
    try:
        # Preprocess the image
        processed_img = preprocess_image_for_ocr(image_path)
        if not processed_img:
            return "", "Preprocessing failed"
        
        # Try different OCR configurations
        configs = [
            '--psm 6',  # Uniform block of text (good for code blocks)
            '--psm 4',  # Single column of text
            '--psm 3',  # Fully automatic
            '--psm 1',  # Automatic with OSD
        ]
        
        best_text = ""
        best_config = ""
        max_length = 0
        
        for config in configs:
            try:
                text = pytesseract.image_to_string(processed_img, config=config)
                if len(text.strip()) > max_length:
                    max_length = len(text.strip())
                    best_text = text
                    best_config = config
            except Exception:
                continue
        
        # Clean up the text
        if best_text:
            # Remove excessive whitespace while preserving structure
            lines = best_text.split('\n')
            cleaned_lines = []
            for line in lines:
                cleaned_line = ' '.join(line.split())  # Clean multiple spaces
                cleaned_lines.append(cleaned_line)
            
            # Remove empty lines at start and end
            while cleaned_lines and not cleaned_lines[0]:
                cleaned_lines.pop(0)
            while cleaned_lines and not cleaned_lines[-1]:
                cleaned_lines.pop()
            
            cleaned_text = '\n'.join(cleaned_lines)
            return cleaned_text, best_config
        
        return "", "No text extracted"
        
    except Exception as e:
        return "", f"OCR error: {e}"

def extract_images_from_docx(docx_path, temp_dir):
    """Extract all images from a .docx file."""
    images = []
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            for file_info in docx_zip.filelist:
                if file_info.filename.startswith('word/media/') and any(
                    file_info.filename.lower().endswith(ext) 
                    for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
                ):
                    image_name = os.path.basename(file_info.filename)
                    temp_image_path = os.path.join(temp_dir, image_name)
                    
                    with docx_zip.open(file_info.filename) as source_file:
                        with open(temp_image_path, 'wb') as target_file:
                            target_file.write(source_file.read())
                    
                    images.append(temp_image_path)
        
        return images
    except Exception as e:
        print(f"    Error extracting images: {e}")
        return []

def add_extracted_text_to_doc(doc, image_name, extracted_text, ocr_method):
    """Add extracted text to the document."""
    if not extracted_text.strip():
        return
    
    # Add header for extracted content
    header_p = doc.add_paragraph()
    header_run = header_p.add_run(f"\n[Extracted Text from {image_name}]")
    header_run.bold = True
    header_run.italic = True
    
    # Add method info
    method_p = doc.add_paragraph()
    method_run = method_p.add_run(f"OCR Method: {ocr_method}")
    method_run.italic = True
    
    # Add the extracted text
    text_p = doc.add_paragraph()
    text_run = text_p.add_run(extracted_text)
    
    # Format as code if it looks like code
    if is_code_like(extracted_text):
        try:
            text_run.font.name = 'Courier New'
            text_p.style = 'Normal'  # Use normal style for code
        except:
            pass
    
    # Add separator
    doc.add_paragraph("-" * 50)

def is_code_like(text):
    """Check if text looks like code."""
    if not text or len(text.strip()) < 20:
        return False
    
    # Programming indicators
    code_patterns = [
        'def ', 'class ', 'import ', 'from ', 'if ', 'else:', 'elif ',
        'for ', 'while ', 'try:', 'except:', 'finally:', 'with ',
        'return ', 'yield ', 'lambda ', 'async ', 'await ',
        '__init__', '__main__', 'self.', 'print(', 'len(',
        '==', '!=', '<=', '>=', '->', '=>', '{}', '[]',
        '    ', '\t',  # Indentation
    ]
    
    # Count indicators
    indicator_count = sum(1 for pattern in code_patterns if pattern in text.lower())
    
    # Check for structural patterns
    has_brackets = ('(' in text and ')' in text) or ('{' in text and '}' in text)
    has_indentation = '    ' in text or '\t' in text
    has_operators = any(op in text for op in ['==', '!=', '<=', '>=', '+=', '-='])
    
    return indicator_count >= 3 or (indicator_count >= 1 and has_brackets and (has_indentation or has_operators))

def process_single_docx_file(input_path, output_path):
    """Process a single .docx file."""
    print(f"Processing: {input_path.name}")
    
    # Create temporary directory
    temp_dir = f"temp_{input_path.stem}"
    os.makedirs(temp_dir, exist_ok=True)
    
    try:
        # Extract images
        images = extract_images_from_docx(input_path, temp_dir)
        
        if not images:
            # No images, just copy the file
            shutil.copy2(input_path, output_path)
            print(f"  No images found - copied as-is")
            return
        
        print(f"  Found {len(images)} images")
        
        # Load the document
        doc = Document(input_path)
        
        # Track if any text was extracted
        extracted_any = False
        
        # Process each image
        for i, image_path in enumerate(images):
            image_name = os.path.basename(image_path)
            print(f"    Processing {i+1}/{len(images)}: {image_name}")
            
            # Extract text using OCR
            extracted_text, ocr_method = extract_text_with_ocr(image_path)
            
            if extracted_text.strip():
                print(f"      Extracted {len(extracted_text)} characters")
                add_extracted_text_to_doc(doc, image_name, extracted_text, ocr_method)
                extracted_any = True
                
                # If it looks like code, note it
                if is_code_like(extracted_text):
                    print(f"      → Looks like CODE!")
            else:
                print(f"      No text extracted ({ocr_method})")
        
        # Save the document
        doc.save(output_path)
        
        if extracted_any:
            print(f"  ✓ Saved with extracted text to: {output_path}")
        else:
            print(f"  ✓ Saved (no text extracted) to: {output_path}")
        
    except Exception as e:
        print(f"  ✗ Error processing {input_path}: {e}")
        # Copy original file as fallback
        try:
            shutil.copy2(input_path, output_path)
            print(f"  → Copied original file as fallback")
        except Exception as copy_error:
            print(f"  → Failed to copy original file: {copy_error}")
    
    finally:
        # Clean up temporary directory
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

def main():
    """Main processing function."""
    book_dir = Path("book")
    output_dir = Path("book_1")
    
    if not book_dir.exists():
        print("Error: 'book' directory not found!")
        return
    
    # Create output directory
    output_dir.mkdir(exist_ok=True)
    
    # Find all .docx files
    docx_files = list(book_dir.glob("*.docx"))
    
    if not docx_files:
        print("No .docx files found!")
        return
    
    print(f"Found {len(docx_files)} .docx files to process")
    print("=" * 60)
    
    # Process each file
    for i, docx_file in enumerate(docx_files, 1):
        print(f"\n[{i}/{len(docx_files)}]", end=" ")
        output_file = output_dir / docx_file.name
        process_single_docx_file(docx_file, output_file)
    
    print("\n" + "=" * 60)
    print("Processing complete!")
    print(f"All processed files are in the '{output_dir}' directory.")
    print("\nFiles with extracted text will have additional content appended")
    print("after the original document content.")

if __name__ == "__main__":
    main()